import unittest
import tempfile
import os
try:
    from docx import Document
except Exception:
    Document = None
try:
    from app.core.word_processor import WordProcessor
except Exception:
    WordProcessor = None


@unittest.skipUnless(Document is not None and WordProcessor is not None, "python-docx or WordProcessor not installed")
class TestWordProcessor(unittest.TestCase):
    def build_temp_docx(self, runs):
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc = Document()
        para = doc.add_paragraph()
        for text in runs:
            para.add_run(text)
        doc.save(path)
        return path

    def test_extract_placeholders(self):
        path = self.build_temp_docx(["{{a}} ", "{{b}}"])
        wp = WordProcessor(path)
        self.assertEqual(sorted(wp.extract_placeholders()), ["a", "b"])
        os.remove(path)

    def test_fill_placeholders_split_run(self):
        path = self.build_temp_docx(["Hello ", "{{tem", "perature_reception", "}} world"])
        out_path = path.replace(".docx", "_out.docx")
        wp = WordProcessor(path)
        wp.fill_placeholders({"{{temperature_reception}}": "4°C"}, out_path)

        result = Document(out_path)
        self.assertNotIn("{{temperature_reception}}", "".join(p.text for p in result.paragraphs))
        self.assertIn("4°C", "".join(p.text for p in result.paragraphs))
        os.remove(path)
        os.remove(out_path)


if __name__ == "__main__":
    unittest.main()
