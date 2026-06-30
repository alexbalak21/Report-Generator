import datetime
import re
import unittest
import tempfile
import os
try:
    from openpyxl import Workbook
    from docx import Document
except Exception:
    Workbook = None
    Document = None
try:
    from app.core.report_generator import ReportGenerator
    from app.core.mapping_loader import MappingLoader
    from app.core.word_processor import WordProcessor
    from app.core.excel_reader import ExcelReader
    import app.repository.config_repository as config_repository
except Exception:
    ReportGenerator = None
    MappingLoader = None
    WordProcessor = None
    ExcelReader = None
    config_repository = None
import json


class TestReportGenerator(unittest.TestCase):
    def setUp(self):
        if Workbook is None or Document is None or ReportGenerator is None:
            self.skipTest("openpyxl, python-docx, or ReportGenerator dependencies not installed")
        self.temp_dir = tempfile.mkdtemp()
        self.xlsx_path = os.path.join(self.temp_dir, "data.xlsx")
        self.template_path = os.path.join(self.temp_dir, "template.docx")
        self.mapping_path = os.path.join(self.temp_dir, "mapping.json")
        self.output_path = os.path.join(self.temp_dir, "output.docx")

        # Create Excel workbook with one data row.
        # K-value is stored as a formula cell so the reader must resolve its computed value.
        wb = Workbook()
        ws = wb.active
        ws.append(["date rapport", "temperature reception", "imp", "k value", "date emabalage", "dlc", "numero echantillon"])
        ws.append([
            datetime.date(2026, 6, 21),
            "4°C",
            0.738026819923372,
            None,
            datetime.date(2026, 5, 31),
            datetime.date(2026, 6, 14),
            1,
        ])
        ws["D2"] = "=1-C2"
        wb.save(self.xlsx_path)

        # Create Word template with placeholders
        doc = Document()
        doc.add_paragraph("Date de rapport : {{date_rapport}}")
        doc.add_paragraph("Température : {{temperature_reception}}")
        doc.add_paragraph("IMP : {{imp}}")
        doc.add_paragraph("K-value : {{k_value}}")
        doc.add_paragraph("Emballage : {{date_emabalage}}")
        doc.add_paragraph("DLC : {{dlc}}")
        doc.add_paragraph("Report : {{numéro_rapport}}")
        doc.save(self.template_path)

        mapping = {
            "config": {
                "data_file": self.xlsx_path,
                "template_file": self.template_path,
                "output_dir": self.temp_dir,
                "date_format": "%d/%m/%Y"
            },
            "date rapport": {"column": "date rapport", "placeholder": "{{date_rapport}}"},
            "temperature reception": {"column": "temperature reception", "placeholder": "{{temperature_reception}}"},
            "imp": {
                "column": "imp",
                "placeholder": "{{imp}}",
                "operations": [
                    {"type": "formula"},
                    {"type": "multiply", "value": 100},
                    {"type": "round", "decimals": 0},
                    {"type": "suffix", "value": "%"}
                ]
            },
            "k value": {
                "column": "k value",
                "placeholder": "{{k_value}}",
                "operations": [
                    {"type": "formula"},
                    {"type": "multiply", "value": 100},
                    {"type": "round", "decimals": 0},
                    {"type": "suffix", "value": "%"}
                ]
            },
            "date emabalage": {
                "column": "date emabalage",
                "placeholder": "{{date_emabalage}}",
                "operations": [{"type": "date_format", "format": "%d/%m/%Y"}]
            },
            "dlc": {
                "column": "dlc",
                "placeholder": "{{dlc}}",
                "operations": [{"type": "date_format", "format": "%d/%m/%Y"}]
            },
            "numero echantillon": {
                "column": "numero echantillon",
                "placeholder": "{{numero_echantillon}}"
            },
            "numero_rapport": {
                "operation": "excel_day_counter",
                "date_column": "date rapport",
                "sample_column": "numero echantillon",
                "date_format": "%d/%m/%Y"
            }
        }
        with open(self.mapping_path, "w", encoding="utf-8") as f:
            json.dump(mapping, f, indent=2, ensure_ascii=False)

    def tearDown(self):
        import shutil

        if os.path.isdir(self.temp_dir):
            shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_report_generation_fills_placeholders(self):
        rg = ReportGenerator(self.xlsx_path, self.template_path, self.mapping_path)
        output = rg.generate(2, self.output_path)
        self.assertEqual(output, self.output_path)

        doc = Document(output)
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("{{date_rapport}}", text)
        self.assertNotIn("{{temperature_reception}}", text)
        self.assertNotIn("{{imp}}", text)
        self.assertNotIn("{{k_value}}", text)
        self.assertNotIn("{{date_emabalage}}", text)
        self.assertNotIn("{{dlc}}", text)
        self.assertNotIn("{{numéro_rapport}}", text)

        self.assertIn("21/06/2026", text)
        self.assertIn("4°C", text)
        self.assertIn("74%", text)
        self.assertIn("26%", text)
        self.assertIn("31/05/2026", text)
        self.assertIn("14/06/2026", text)
        self.assertIn("260621-1", text)

    def test_report_generation_with_xlsx_mapping(self):
        if MappingLoader is None:
            self.skipTest("MappingLoader unavailable")

        mapping_xlsx_path = os.path.join(self.temp_dir, "mapping.xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = "mappings"
        ws.append(["Spreadsheet Column", "Placeholder", "Operation", "Type", "Notes"])
        ws.append(["date rapport", "{{date_rapport}}", "", "date (input)", ""])
        ws.append(["temperature reception", "{{temperature_reception}}", "", "text (input)", ""])
        ws.append(["date emabalage", "{{date_emabalage}}", "date_format %d/%m/%Y", "date", ""])
        ws.append(["dlc", "{{dlc}}", "date_format %d/%m/%Y", "date", ""])
        ws.append(["numero echantillon", "{{numero_echantillon}}", "", "text/number (input)", ""])
        ws.append(["(computed) numero_rapport", "—", "report_number (date_column: date rapport; sample_column: numero echantillon)", "computed", ""])
        ws.append(["(computed) file_name", "—", "format(\"{name} {numero_rapport}.docx\")", "computed", ""])
        wb.save(mapping_xlsx_path)

        loader = MappingLoader(mapping_xlsx_path)
        loader.update_file_name_field(name="TEST")
        loaded = loader.load()
        self.assertEqual(loaded["file_name"]["name"], "TEST")
        self.assertEqual(loaded["file_name"]["operation"], "format")

        output = ReportGenerator(self.xlsx_path, self.template_path, mapping_xlsx_path).generate(2, self.output_path)
        self.assertEqual(output, self.output_path)

        doc = Document(output)
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("21/06/2026", text)
        self.assertIn("4°C", text)
        self.assertIn("260621-1", text)

    def test_load_file_name_name_from_xlsx_format_rule(self):
        if MappingLoader is None:
            self.skipTest("MappingLoader unavailable")

        mapping_xlsx_path = os.path.join(self.temp_dir, "mapping.xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = "mappings"
        ws.append(["Spreadsheet Column", "Placeholder", "Operation", "Type", "Notes"])
        ws.append(["(computed) file_name", "—", "format(\"NOVOCIB Rapport d'essai {numero_rapport}.docx\")", "computed", ""])
        wb.save(mapping_xlsx_path)

        loader = MappingLoader(mapping_xlsx_path)
        fn = loader.load_file_name_field()
        self.assertEqual(fn.get("name"), "NOVOCIB Rapport d'essai {numero_rapport}")

    def test_load_config_from_xlsx_config_sheet_with_header(self):
        if MappingLoader is None:
            self.skipTest("MappingLoader unavailable")

        mapping_xlsx_path = os.path.join(self.temp_dir, "mapping.xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = "mappings"
        ws.append(["Spreadsheet Column", "Placeholder", "Operation", "Type", "Notes"])
        ws.append(["date rapport", "{{date_rapport}}", "", "date (input)", ""])
        ws.append(["(computed) numero_rapport", "—", "report_number (date_column: date rapport; sample_column: numero echantillon)", "computed", ""])

        config = wb.create_sheet("config")
        config.append(["Key", "Value"])
        config.append(["data_file", self.xlsx_path])
        config.append(["template_file", self.template_path])
        config.append(["output_dir", self.temp_dir])
        config.append(["date_format", "%d/%m/%Y"])
        wb.save(mapping_xlsx_path)

        loader = MappingLoader(mapping_xlsx_path)
        cfg = loader.load_config()

        self.assertEqual(cfg, {
            "data_file": self.xlsx_path,
            "template_file": self.template_path,
            "output_dir": self.temp_dir,
            "date_format": "%d/%m/%Y",
        })

    def test_mapping_registry_persists_mapping_path(self):
        if config_repository is None:
            self.skipTest("config_repository unavailable")

        config_repository.DB_PATH = os.path.join(self.temp_dir, "app_data.db")
        mapping_path = os.path.join(self.temp_dir, "mapping.xlsx")
        with open(mapping_path, "w", encoding="utf-8") as f:
            f.write("dummy")

        first_id = config_repository.mapping_add(mapping_path)
        self.assertGreater(first_id, 0)

        second_id = config_repository.mapping_add(mapping_path)
        self.assertEqual(first_id, second_id)

        items = config_repository.mapping_list()
        self.assertEqual(items, [(first_id, mapping_path)])

    def test_last_mapping_placeholder_column_comparison(self):
        if config_repository is None or WordProcessor is None or ExcelReader is None:
            self.skipTest("Required components unavailable")

        config_repository.DB_PATH = os.path.join(self.temp_dir, "app_data.db")
        mapping_path = os.path.join(self.temp_dir, "mapping.xlsx")

        wb = Workbook()
        ws = wb.active
        ws.title = "mappings"
        ws.append(["Spreadsheet Column", "Placeholder", "Operation", "Type", "Notes"])
        ws.append(["date rapport", "{{date_rapport}}", "", "date (input)", ""])
        ws.append(["temperature reception", "{{temperature_reception}}", "", "text (input)", ""])
        ws.append(["numero rapport", "{{numero_rapport}}", "", "text (input)", ""])
        ws.append(["numero echantillon", "{{numero_echantillon}}", "", "text (input)", ""])

        config = wb.create_sheet("config")
        config.append(["Key", "Value"])
        config.append(["data_file", self.xlsx_path])
        config.append(["template_file", self.template_path])
        config.append(["output_dir", self.temp_dir])
        config.append(["date_format", "%d/%m/%Y"])
        wb.save(mapping_path)

        config_repository.mapping_add(mapping_path)
        last_mapping = config_repository.mapping_get_last()
        self.assertEqual(last_mapping, mapping_path)

        word = WordProcessor(self.template_path)
        placeholders = set(word.extract_placeholders())

        excel = ExcelReader(self.xlsx_path)
        excel.load()
        columns = {col.strip() for col in excel.get_columns()}

        mapping = MappingLoader(last_mapping).load()
        mapping_placeholders = set()
        mapping_columns = set()
        for key, rule in mapping.items():
            if isinstance(rule, dict):
                if "column" in rule:
                    mapping_columns.add(str(rule.get("column", "")).strip())
                placeholder = rule.get("placeholder")
                if isinstance(placeholder, str):
                    match = re.search(r"\{\{(.*?)\}\}", placeholder)
                    if match:
                        mapping_placeholders.add(match.group(1).strip())
                elif key not in {"file_name"}:
                    mapping_placeholders.add(key)

        missing_placeholders = placeholders - mapping_placeholders
        missing_columns = mapping_columns - columns
        extra_placeholders = mapping_placeholders - placeholders

        self.assertFalse(
            missing_placeholders,
            f"Missing placeholder mappings for template placeholders: {sorted(missing_placeholders)}"
        )
        self.assertFalse(
            missing_columns,
            f"Mapped Excel columns missing from spreadsheet: {sorted(missing_columns)}"
        )
        self.assertFalse(
            extra_placeholders,
            f"Mapped placeholders not found in Word template: {sorted(extra_placeholders)}"
        )


if __name__ == "__main__":
    unittest.main()
