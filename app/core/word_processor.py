import re

try:
    from docx import Document
    from docx.oxml.ns import qn
except Exception:
    Document = None
    qn = None

PLACEHOLDER_RE = re.compile(r"\{\{.*?\}\}")


class WordProcessor:
    def __init__(self, filepath):
        if Document is None:
            raise ImportError("python-docx is required for WordProcessor")
        self.filepath = filepath
        self.document = Document(filepath)

    # ------------------------------------------------------------------
    # Iterate every paragraph in the document including tables
    # ------------------------------------------------------------------

    def _all_paragraphs(self):
        yield from self.document.paragraphs
        for table in self.document.tables:
            yield from self._paragraphs_in_table(table)

    def _paragraphs_in_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    yield from self._paragraphs_in_table(nested)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def extract_placeholders(self) -> list:
        found = []
        for para in self._all_paragraphs():
            found.extend(PLACEHOLDER_RE.findall(para.text))
        return [placeholder.strip("{}") for placeholder in found]

    def fill_placeholders(self, mapping: dict, output_path: str) -> None:
        for para in self._all_paragraphs():
            self._replace_in_paragraph(para, mapping)
            if self._has_remaining_placeholder(para.text):
                self._replace_in_paragraph(para, mapping)
        self.document.save(output_path)

    # ------------------------------------------------------------------
    # Core replacement — preserves per-run formatting
    # ------------------------------------------------------------------

    def _replace_in_paragraph(self, paragraph, mapping: dict) -> None:
        if not paragraph.runs:
            return

        if not self._has_remaining_placeholder("".join(r.text for r in paragraph.runs)):
            return

        self._merge_split_placeholders(paragraph)
        self._replace_runs(paragraph.runs, mapping)

    def _replace_runs(self, runs, mapping: dict) -> None:
        for run in runs:
            if not self._has_remaining_placeholder(run.text):
                continue
            run.text = self._replace_run_text(run.text, mapping)

    def _replace_run_text(self, text: str, mapping: dict) -> str:
        for placeholder, value in mapping.items():
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        return text

    @staticmethod
    def _has_remaining_placeholder(text: str) -> bool:
        return "{{" in text and "}}" in text

    def _merge_split_placeholders(self, paragraph) -> None:
        """
        Merge runs so that every {{...}} placeholder lands in a single run.

        Word sometimes splits a placeholder like {{qualite}} across many runs,
        even splitting the opening {{ itself into ' {' + '{'. A simple
        search for '{{' inside individual runs misses that. Instead we:
          1. Join all run texts to get the full paragraph text.
          2. Use the regex to find every placeholder and its char span.
          3. Figure out which runs each placeholder spans and merge them.
          4. Process matches in reverse order so earlier indices stay valid.
        """
        runs = paragraph.runs
        if len(runs) <= 1:
            return

        full_text = "".join(r.text for r in runs)
        matches = list(PLACEHOLDER_RE.finditer(full_text))
        if not matches:
            return

        # Starting character position of each run inside full_text
        starts = []
        pos = 0
        for run in runs:
            starts.append(pos)
            pos += len(run.text)

        def run_at(char_pos: int) -> int:
            for idx in range(len(runs) - 1, -1, -1):
                if starts[idx] <= char_pos:
                    return idx
            return 0

        # Process in reverse so merging higher-index runs does not shift lower ones
        for m in reversed(matches):
            first = run_at(m.start())
            last  = run_at(m.end() - 1)
            if first == last:
                continue
            merged = "".join(r.text for r in runs[first : last + 1])
            runs[first].text = merged
            for k in range(first + 1, last + 1):
                runs[k].text = ""
